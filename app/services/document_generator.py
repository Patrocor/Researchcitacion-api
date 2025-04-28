from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from fastapi.responses import StreamingResponse
from io import BytesIO
from app.models.response_model import Paragraph
import docx.oxml

def generate_docx(sections: list[dict], references: list[dict], topic: str) -> StreamingResponse:
    """
    Genera un archivo Word (.docx) basado en las secciones, párrafos y referencias proporcionadas.
    """

    doc = Document()

    # Insertar número de página en el pie de página
    section = doc.sections[0]
    footer = section.footer
    paragraph_footer = footer.paragraphs[0]
    paragraph_footer.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    paragraph_footer.add_run("Página ").font.size = Pt(10)
    field_code = 'PAGE  \\* MERGEFORMAT'
    paragraph_footer._p.append(docx.oxml.parse_xml(f'<w:fldSimple w:instr="{field_code}" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'))

    # (luego sigue igual todo lo demás como ya tenías)

    # Título principal, párrafos, referencias, guardar el archivo, etc.